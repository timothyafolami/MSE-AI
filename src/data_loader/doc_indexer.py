import os
import sys

import docx
from langchain_community.document_loaders import TextLoader
from langchain.schema import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from src.data_loader import settings
from loguru import logger
from src.data_loader.pdf_loader import load_pdf 

def create_and_save_document_index(embeddings, document_path):
    """
    Processes documents (PDF, DOCX/DOC, TXT), creates embeddings, and saves to a single FAISS index.
    
    Args:
        embeddings: The embeddings object to use
        document_path (str): Path to the document
        
    Returns:
        str: Path where the index was saved
    """
    # Ensure output directory exists
    os.makedirs(settings.DOC_INDEXES_DIR, exist_ok=True)
    
    # Use a common index name for all documents
    save_path = os.path.join(settings.DOC_INDEXES_DIR, "materials_database")
    
    # Load document based on file type
    file_ext = os.path.splitext(document_path)[1].lower()
    doc_name = os.path.splitext(os.path.basename(document_path))[0]
    
    try:
        if file_ext == '.pdf':
            # Use the load_pdf function from pdf_loader.py
            documents = load_pdf(document_path)
        elif file_ext in ['.docx', '.doc']:
            doc = docx.Document(document_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text = "\n\n".join(paragraphs)
            documents = [Document(page_content=text, metadata={"source": document_path, "type": "word", "doc_name": doc_name})]
        elif file_ext == '.txt':
            loader = TextLoader(document_path)
            documents = loader.load()
            # Add doc_name to metadata
            for doc in documents:
                doc.metadata["doc_name"] = doc_name
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    except Exception as e:
        logger.error(f"Error loading document {document_path}: {str(e)}")
        raise ValueError(f"Failed to load or process document: {document_path} due to {str(e)}")

    # Split text into chunks
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=2000,
        chunk_overlap=150,  # Increased overlap for better context preservation
        length_function=len,
        separators=["\n\n", "\n", ". ", " ", ""]
    )
    split_documents = text_splitter.split_documents(documents)
    
    # Add document name to metadata if not already present
    for doc in split_documents:
        if "doc_name" not in doc.metadata:
            doc.metadata["doc_name"] = doc_name
    
    logger.info(f"Split document {doc_name} into {len(split_documents)} chunks")

    try:
        # Create a proper HuggingFaceEmbeddings wrapper around SentenceTransformer
        # This is needed because SentenceTransformer uses .encode() while LangChain expects .embed_documents()
        from langchain_community.embeddings import HuggingFaceEmbeddings
        
        # Create HuggingFace embeddings from the model
        hf_embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
        
        # Check if index already exists
        if os.path.exists(os.path.join(save_path, "index.faiss")):
            # Load existing index
            logger.info(f"Loading existing index at {save_path}")
            vectorstore = FAISS.load_local(save_path, hf_embeddings, allow_dangerous_deserialization=True)
            
            # Add new documents to existing index
            vectorstore.add_documents(split_documents)
        else:
            # Create new vector store from split documents
            logger.info(f"Creating new index at {save_path}")
            vectorstore = FAISS.from_documents(split_documents, hf_embeddings)
        
        # Ensure directory exists for saving the index
        os.makedirs(os.path.dirname(save_path), exist_ok=True)
        
        # Save the FAISS index
        vectorstore.save_local(save_path)
    except Exception as e:
        logger.error(f"Error creating or saving FAISS index for {document_path}: {str(e)}")
        raise ValueError(f"Failed to create or save index for document: {document_path} due to {str(e)}")
    
    logger.success(f"Successfully indexed {document_path} → {save_path}")
    return save_path

def retrieve_documents(
    embeddings,
    query: str,
    document_name: str = None,  # Now optional
    search_type: str = "similarity",  # Changed default to similarity
    k: int = 5
) -> list:
    """
    Retrieve relevant documents from the materials database based on a query.
    
    Args:
        embeddings: The embeddings object to use (ignored, we create our own)
        query (str): Search query or question
        document_name (str, optional): Name of the specific document index to search.
                                      If None, searches the unified database.
        search_type (str): Type of search ('mmr' or 'similarity')
        k (int): Number of documents to return
    
    Returns:
        List[str]: Relevant document chunks
    """
    # Use unified index by default, or specific document index if provided
    if document_name:
        index_path = os.path.join(settings.DOC_INDEXES_DIR, document_name)
    else:
        index_path = os.path.join(settings.DOC_INDEXES_DIR, "materials_database")
    
    try:
        # Create a proper HuggingFaceEmbeddings wrapper
        from langchain_community.embeddings import HuggingFaceEmbeddings
        hf_embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
        
        logger.info(f"Loading index from {index_path}")
        vector_store = FAISS.load_local(index_path, hf_embeddings, allow_dangerous_deserialization=True)
        
        # Simple similarity search - most reliable approach
        docs = vector_store.similarity_search(query, k=k)
        index_name = document_name if document_name else "unified materials database"
        logger.success(f"Retrieved {len(docs)} documents from {index_name} for query: {query}")
        return docs
    except Exception as e:
        index_name = document_name if document_name else "unified materials database"
        logger.error(f"Failed to retrieve documents for {index_name}: {str(e)}")
        
        # Return empty list instead of raising an exception
        logger.warning(f"Returning empty results due to retrieval error")
        return []